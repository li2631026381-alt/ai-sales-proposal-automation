# export_docx.py
from docx import Document

def export_to_docx(data: dict, out_path: str) -> None:
    doc = Document()
    doc.add_heading("AI Sales Proposal Automation – Proposal", level=1)

    doc.add_heading("Executive Summary", level=2)
    doc.add_paragraph(data.get("executive_summary", ""))

    doc.add_heading("Pain Point Analysis", level=2)
    doc.add_paragraph(data.get("pain_point_analysis", ""))

    doc.add_heading("Proposed AI Solution", level=2)
    sol = data.get("proposed_ai_solution", {})
    doc.add_paragraph(f"Input: {sol.get('input','')}")
    doc.add_paragraph(f"Processing: {sol.get('processing','')}")
    doc.add_paragraph(f"Output: {sol.get('output','')}")
    doc.add_paragraph("Modules: " + ", ".join(sol.get("modules", [])))

    doc.add_heading("Implementation Plan", level=2)
    for ph in data.get("implementation_plan", []):
        doc.add_paragraph(ph["phase"], style="List Bullet")
        for it in ph["items"]:
            doc.add_paragraph(it, style="List Bullet 2")

    doc.add_heading("ROI Estimation", level=2)
    doc.add_paragraph(data.get("roi_estimation_notes", ""))

    doc.add_heading("Pricing Structure", level=2)
    for t in data.get("pricing_structure", []):
        doc.add_paragraph(f"{t['tier']} ({t['price_note']}): " + "; ".join(t["includes"]), style="List Bullet")

    doc.add_heading("Risks & Mitigations", level=2)
    for r in data.get("risks_mitigations", []):
        doc.add_paragraph(f"{r['risk']} → {r['mitigation']}", style="List Bullet")

    doc.add_heading("Next Steps", level=2)
    for s in data.get("next_steps", []):
        doc.add_paragraph(s, style="List Bullet")

    doc.add_heading("Sales Email Draft", level=2)
    email = data.get("sales_email_draft", {})
    doc.add_paragraph("EN:", style=None)
    doc.add_paragraph(email.get("en", ""))
    doc.add_paragraph("ZH:", style=None)
    doc.add_paragraph(email.get("zh", ""))

    doc.save(out_path)